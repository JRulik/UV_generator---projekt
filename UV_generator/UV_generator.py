import glob
import os, sys
import signal
import time
import docx
from datetime import date
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
#from copy import deepcopy


#PODMINKY
#- nainstalovat python3 - https://www.python.org/downloads/
#- nainstalovat python-docx modul - do cmd dat: pip install python-docx

#-----------------------------------------------------------------------------------------------
#globalni promenne 

#flagy
pkg_flag = 0
pkg_tool_flag = 0
pkg_project_flag = 0
ufd_flag = 0
UVTameplate ='F083_03.docx'

#Verze SW
UV_version="0.50"

#cesta k souboru
path = r"C:\Users\Adminator\Desktop\PROGRAMOVANI_ISC-PROJEKT\203054_ISC-TEST"

#vsechny soubory v prog. sade pro vypis
g_files=[]

#jmeno souboru
#name = None

#Uvodni vypis s verzi
intro = (
        "-------------------- ---------------------- ----------------------\n"
        "--------------------      UV maker " + UV_version + "     ----------------------\n"
        "-------------------- ---------------------- ----------------------\n"
        "cmd - ptyhon UV_maker.py <cestaKProgamoveSade>\n"
        "-------------------- ---------------------- ----------------------\n"
        )
#------------------------------------------------------------------------------------------------


#-------------------------------------------------------------------
#Nalezeni app slozky a manifestu a vlozeni do g_files
#Argument:
#   files - pole souboru z nadrazene slozky 
#-------------------------------------------------------------------
def application_file_add(files):

    #predpoklad ze vsechny aplikace jsou .tar.gz
    for file in files:
        if file.find(".tar.gz")!=-1:
            #najdi manifest
            for soubor in files:
                if soubor.find("manifest.txt")!=-1:
                #otevri manifest a vypis verzi + PID
                    try:
                        file_manifest = open(soubor, "r")
                        for line in file_manifest:
                            if "package.pid" in line:
                                startIndex = line.find("=")+1 #+1 posunuti indexu, +1 mezera pred chtenou cislici
                                PID = line[startIndex:line.find("\\n")]
                            if "package.major" in line:
                                startIndex = line.find("=")+1+1
                                versionMajor = line[startIndex:line.find("\\n")]
                            if "package.minor" in line:
                                startIndex = line.find("=")+1+1
                                versionMinor = line[startIndex:line.find("\\n")]
                            if "package.comment" in line:
                                startIndex = line.find("=")+1+1
                                comment = line[startIndex:line.find("\\n")]
                        file_manifest.close()
                    except:
                        print("Nejde otevrit manifest.txt k souboru: "+file)
                    g_files.append(file+" PID " + PID + ", v" +versionMajor +"."+versionMinor +", comment: "+comment)
        #detekce komprimovaneho souboru OS Linux
        if file.find(".squashfs")!=-1:
            g_files.append(file);

#-------------------------------------------------------------------
#Vypsani pkg a "IscConfig.exe" do g_files
#-------------------------------------------------------------------
def pkg_file_add(files,root_path):
    global pkg_flag
    pkg_flag=1
    #print("-_PKG_-")
    ##vypis jen iscconfig, kdyz ho najdes ->asi zbytecny, mohlo by jit natvrdo                      #TODO - upravit aby ne natvrdo, jsou tam taky iscconfig.exe.config soubory
    #for file in files:
        #if file.find("IscConfig.exe") !=-1 : #nalezeni IscConfig
            #g_files.append(root_path)
            #g_files.append(file)
    g_files.append(root_path)
    g_files.append(os.path.join(root_path, "IscConfig.exe"))

#-------------------------------------------------------------------
#Vypsani konfigurace a firmware ve slozce AXIS-P3905R
#-------------------------------------------------------------------
def camera_file_add(files):
    cmt_cam_flag = 0
    #predpoklad ze vsechny aplikace jsou .tar.gz
    for file in files:
        if cmt_cam_flag == 0 and file.find(".cmt")!=-1:
            cmt_cam_flag=1
            g_files.append(file)                     #TODO zjistik MKI/MKII z nazvu slozky, precist verzi konfigurace z nazvu
        elif file.find(".bin")!=-1:
            g_files.append(file)  


#-------------------------------------------------------------------
#Prochazeni souboru a slozek na zadane ceste. Soubory a slozky jsou rozdeleny do dvou poli.
#Nejdriv vypise soubory ve slozce, nasledne projde prvni nalezenou slozku
#Argument:
#   root_path - cesta k slozce, ve ktere se vyhledava
#-------------------------------------------------------------------
def files_and_subdirectories(root_path):
    global pkg_tool_flag, pkg_flag, pkg_project_flag, cmt_cam_flag, ufd_flag
    PID,versionMajor,versionMinor,comment = "NaN","NaN","NaN","NaN"
    files = []
    directories = []

    #----------------------------------------- Rozdeleni souboru a slozek v root slozce --------------------------------------------
    #-----------------------------------------------------------------------------------------------------------------

    for f in os.listdir(root_path):
        f=os.path.join(root_path, f) #doplneni komplet cesty k souboru
        #je to soubor
        if os.path.isfile(f):
            files.append(f)
        #je to slozka
        elif os.path.isdir(f):
            #files_and_subdirectories(f)    #prohle do hloubky ->upravit pak i zbytek metody
            directories.append(f)

    #----------------------------------------- PRIDANI PRO VYPIS -----------------------------------------------------
    #-----------------------------------------------------------------------------------------------------------------

    #filtr souboru v aplikacni slozce MS,AIPD,AIPQ,AIPM,APB,PPMB -> najde .tar.gz a vypise k nemu manifest
    # jsme v _pkg_ a slozce Project
    #detekce MS,AIPD,AIPM,AIPQ,APB,PPMB - hledani app i v podslozek, projeti manifestu pro verzi
    if root_path.find("_pkg_")!=-1 and root_path.find("Project")!=-1 and (root_path.find("AIPM")!=-1 or root_path.find("MS")!=-1 or root_path.find("AIPD")!=-1 or root_path.find("AIPQ")!=-1 or root_path.find("APB")!=-1 or root_path.find("PPMB")!=-1):
        application_file_add(files)
    #detekce _pkg_ slozky <- pro trignuti musi root byt _pkg_ slozka
    elif pkg_flag == 0 and root_path.find("_pkg_")!=-1: 
        pkg_file_add(files,root_path)
    #detekce Tools v _pkg_ slozce + kontrola flagu <- nic se nepridava do prohledani/ze osuboru
    elif pkg_tool_flag == 0  and root_path.find("_pkg_")!=-1 and root_path.find("Tools")!=-1 :
        pkg_tool_flag = 1
        directories.clear()
    #detekce slozky Project ve slozce _pkg_ -> pro nevypsani definic projektu .xml -> if jenom aby neproclo konecny else
    elif pkg_project_flag == 0 and root_path.find("_pkg_")!=-1 and root_path.find("Project")!=-1:
        pkg_project_flag=1
    #detekce slozky AXIS-P3905Rxxxxxxx ve Project ve slozce _pkg_ -> detekce firmware a 1 konfigurace                                   #TODO: doplnit vic druhu kamer
    elif  root_path.find("_pkg_")!=-1 and root_path.find("Project")!=-1  and root_path.find("AXIS-P3905R")!=-1:
          camera_file_add(files)  
    #detekce UFD config slozky
    elif  ufd_flag == 0 and root_path.find("ufd-config")!=-1:
        ufd_flag = 1 
        g_files.append(os.path.join(root_path, "UsbFlashDiskConfigurator.exe"))
    else:
    #vypis souboru/pridani souboru do glob pole vsech souboru v sade
        for file in files:
            #print(file)
            g_files.append(file)

    #rekurzivni volani na prvni slozku v subdir
    for subdir in directories:
        files_and_subdirectories(subdir)

    #useless return
    return directories, files


#-------------------------------------------------------------------
#detekce spusteni z konzole s argumenty/samostatne
#Argument:
#   1 - cestka k programove sade
#-------------------------------------------------------------------
def inputDetection():
    global path,name
    if len(sys.argv) == 1:
        print (intro)
        path = input("Cesta k programove sade: ")                    # --odkomentovat--------------
        #name = input("Název protokolu: ") 
    else:
        #print ('Argument List:', str(sys.argv))
        print (intro)
        path = sys.argv[1]                                         # --odkomentovat-----------------
        print ('Cesta k programove sade: ',path)
        #name = sys.argv[2]

#-------------------------------------------------------------------
#Koporovani tabulky za urcity paragraf
#Argument:
#   1 - tabulka ke zkopirovani
#   2 - paragraf za ktery kopirovat
#-------------------------------------------------------------------
def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)



#-------------------------------------------------------------------------------------------------------------------------------------
#------------------------------- -------------------------------  MAIN ------------------------------- -------------------------------
#-------------------------------------------------------------------------------------------------------------------------------------

today = date.today()
#print ('Number of arguments:', len(sys.argv), 'arguments.')

#detekce spusteni z cmd/samostatne 
inputDetection()

#prohledani slozek, podslozek a souboru na zadane ceste
try:
    files_and_subdirectories(path)
except:
    print("[ERROR] - Asi spatne zadana cesta.")
    input("Press Enter to continue...")
    sys.exit(1)

#zapis nalezenych souboru do dokumentu
try:
    doc = docx.Document(UVTameplate)

    #//TODO WORKSPACE ---------------------------------------------------------------- 
    #copy_table_after(doc.tables[5], doc.paragraphs[13])
    #//TODO WORKSPACE ---------------------------------------------------------------- 

    #vyplneni dokumentu
    counter =0
    for table in doc.tables:                                                            #TODO:// - kopirovani tabulky i s paragrafy nad nima (cislovani, lajna)
        #print ("---------------", counter ,"-----------------")
        #print(len(table.rows))
        #print(len(table.columns))

        #Vypisovani obsahu programove sady

        #najde spravnou tabulku podle poctu radku a sloupcu  #TODO: -> predelat na vyhledavani podle tagu!!!
        if (len(table.rows)==4) and (len(table.columns)==2):
            if len(g_files) > counter:
                
                #najde paragraf pred tabulkou s tagem <SOFTWARE> a vyplniho nazvem SW
                for para in doc.paragraphs:
                    if para.text.find("<SOFTWARE>")!=-1:
                        para.clear()
                        fileStartIndex =  g_files[counter].rfind("\\")
                        #protoze se pridaval PID a verze pro TFT,PPMB apod. do nazvu souboru, musi se zas vyfiltrovat
                        if g_files[counter].find("PID") !=-1:
                            fileEndIndex = g_files[counter].find("PID")
                            run = para.add_run(g_files[counter][fileStartIndex+1:fileEndIndex])
                        else:
                            run = para.add_run(g_files[counter][fileStartIndex+1:])

                        run.font.name='Arial'
                        run.font.bold=True
                        run.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
                        break   
                
                #vyplni polozku "Typ softwaru" nazvem SW
                paragraph = table.cell(0,1).paragraphs[0].clear()
                fileStartIndex =  g_files[counter].rfind("\\")
                run = paragraph.add_run(g_files[counter][fileStartIndex+1:])
                run.font.name='Arial'

               #vyplni polozku "Navod na sestaveni"                                                 #TODO:// - prjit slozku NS -> otevrit pdf -> najit programovou sadu ->pokdu najde ->vypis pdf
                paragraph = table.cell(1,1).paragraphs[0].clear()
                run = paragraph.add_run("N/A")                      
                run.font.name='Arial'

                #vyplneni zdrojovych souboru a pripadne prepsani polozky "Typ software"
                #filtr na binarku pro MK II -> vypsani zdrojovy cesty
                if g_files[counter].find("P3905-R_") != -1 and g_files[counter].find(".bin") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    if g_files[counter].find("Mk"):
                        run = paragraph.add_run("Firmware pro IP kamery AP3905-R MKII verze <DOPLNIT>")
                        run.font.name='Arial'
                        #vyplni polozku zdroj
                        zdroj = "\\Vyvoj\\ACASYS\\KAMERY\\AXIS\\FW\\P3905-R MkII\\<DOPLNIT>" #TODO:// - vypsat verzi podle verze souboru nebo natvrdo
                        paragraph = table.cell(2,1).paragraphs[0].clear()
                        run = paragraph.add_run(zdroj)                      
                        run.font.name='Arial'
                    else:
                        run = paragraph.add_run("Firmware pro IP kamery AP3905-R  verze <DOPLNIT>")
                        run.font.name='Arial'
                        #vyplni polozku zdroj
                        zdroj = "\\Vyvoj\\ACASYS\\KAMERY\\AXIS\\FW\\P3905-R\\<DOOPLNIT>" #TODO:// - vypsat verzi podle verze souboru nbo natvrdo
                        paragraph = table.cell(2,1).paragraphs[0].clear()
                        run = paragraph.add_run(zdroj)                      
                        run.font.name='Arial'

                 #filtr na konfiguraci kamery P3905-R (filtr podle .cmt a slozky po ceste)
                if g_files[counter].find(".cmt") != -1 and g_files[counter].find("P3905") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    run = paragraph.add_run("Konfigurace pro IP kamery AP3905-R verze <DOPLNIT>")
                    run.font.name='Arial'

                #filtr na UFD 101 -> vypsani zdrojovy cesty
                if g_files[counter].find("UsbFlashDiskConfigurator.exe") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    run = paragraph.add_run("Nástroj pro vytvoření bootovací klíčenky verze 1.0.1")
                    run.font.name='Arial'
                    #vyplni polozku zdroj
                    zdroj = "\\Vyvoj\\ACASYS\\SW\\USB Flash Disk Configurator\\DSK\\1.0.1\\" #TODO vypsat vverzi UFD z nazvu slozky
                    paragraph = table.cell(2,1).paragraphs[0].clear()
                    run = paragraph.add_run(zdroj)                      
                    run.font.name='Arial'

                #filtr na _pkg_ pro vypsani polozky "Typ software"
                if g_files[counter][fileStartIndex+1:].find("_pkg_") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    run = paragraph.add_run("Programové vybavení komponent")
                    run.font.name='Arial'

                
                #filtr na ISCconfig pro vypsani polozky "Typ software"
                if g_files[counter][fileStartIndex+1:].find("IscConfig.exe") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    run = paragraph.add_run("Servisní nástroj v<DOPLNIT>")
                    run.font.name='Arial'

                 #filtr na Linux OS .squashfs
                if g_files[counter][fileStartIndex+1:].find(".squashfs") != -1:
                    #prepise polozku "Typ softwaru"
                    paragraph = table.cell(0,1).paragraphs[0].clear()
                    fileStartIndex =  g_files[counter].rfind("\\")
                    run = paragraph.add_run("Operační systém pro <DOPLNIT>, v<DOPLNIT>")
                    run.font.name='Arial'

                #vyplni polozku "Vystupni adresar
                paragraph2 = table.cell(3,1).paragraphs[0].clear()
                #deltedCharactersIndex = g_files[counter].find(":")
                run2 = paragraph2.add_run("\\"+g_files[counter][2:fileStartIndex]+"\\") #2 je misto deltedCharactersIndex
                run2.font.name='Arial'
                counter+=1



        #vyplneni uvodni stranky dokumentu                                              #TODO - udelat flag -> aby jen jednou -> ale bacha, je to vic tabulek
        else:
            for y in range (0,len(table.rows),1):
                for x in range (0,len(table.columns),1):
                    if table.cell(y,x).text == "<DATE>":
                       paragraph = table.cell(y,x).paragraphs[0].clear()
                       run = paragraph.add_run(today.strftime("%d.%m.%Y"))
                       #run.font.name='Arial'   #nevim proc zde neni potreba
                       table.cell(y,x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    if table.cell(y,x).text == "<PATH>":
                       deltedCharactersIndex = path.find(":")
                       paragraph = table.cell(y,x).paragraphs[0].clear()
                       run = paragraph.add_run("\\"+path[deltedCharactersIndex+1:])
                       run.font.name='Arial'
                       run.font.bold=True
                       table.cell(y,x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    if table.cell(y,x).text == "<PACKAGE>":
                       packageStartIndex = path.rfind("\\")
                       paragraph = table.cell(y,x).paragraphs[0].clear()
                       run = paragraph.add_run(path[packageStartIndex+1:])
                       run.font.name='Arial'
                       table.cell(y,x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    if table.cell(y,x).text == "<VERSION>":
                       packageStartIndex = path.rfind("\\")
                       versionEndIndex = path.rfind("_")
                       paragraph = table.cell(y,x).paragraphs[0].clear()
                       run = paragraph.add_run(path[packageStartIndex+1:versionEndIndex])
                       run.font.name='Arial'
                       table.cell(y,x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                       table.cell(y,x).alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if table.cell(y,x).text == "<PROJECT>":
                       versionEndIndex = path.rfind("_")
                       paragraph = table.cell(y,x).paragraphs[0].clear()
                       run = paragraph.add_run(path[versionEndIndex+1:])
                       run.font.name='Arial'
                       table.cell(y,x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                       table.cell(y,x).alignment = WD_ALIGN_PARAGRAPH.CENTER
                #print(table.cell(y,x).text)
except:
    print("Asi nejde otevrit vzorovy formular. Nebo taky uplne jina chyba.\n")
    input("Press Enter to continue...")
    sys.exit(1)

try:
    doc.save('UV.docx')
    print("[OK] - UV soubor uspesne vytvoren.\n")
except:
    print("[ERROR] - Asi mas otevreny ten vystupni soubor vole!\n")
    input("Press Enter to continue...")
    sys.exit(1)

input("Press Enter to continue...")

"""
#Endless sleep
counter = 1
print ("---------------", counter ,"-----------------")
while True:
    #print ("---------------", counter ,"-----------------")
    counter +=1
    time.sleep(1)
"""